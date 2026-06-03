from __future__ import annotations

import io
from collections import defaultdict
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .models import ContractContext, MatrixRow
from .scoring import (
    IMPACT_OPTIONS,
    PROBABILITY_OPTIONS,
    canonical_impact,
    canonical_probability,
    risk_score,
    score_value,
)


CATEGORY_TITLES = {
    "planejamento": "1 - Riscos do Planejamento da Contratação",
    "selecao": "2 - Riscos da Seleção de Fornecedor",
    "gestao": "3 - Riscos da Gestão do Contrato",
    "solucao": "4 - Riscos da Solução",
    "instalacao": "4 - Riscos da Instalação dos Equipamentos",
    "cronograma": "5 - Riscos para o cronograma da contratação",
}

LEVEL_OPTIONS = [
    "1 a 3 - Pequeno",
    "4 a 6 - Moderado",
    "8 a 12 - Alto",
    "15 a 25 - Crítico",
]

STRATEGY_OPTIONS = ["Mitigar", "Aceitar", "Compartilhar", "Evitar"]


def to_docx(rows: list[MatrixRow], context: ContractContext) -> bytes:
    document = Document()
    _configure_document(document)

    _add_cover(document, context)
    _add_history(document)
    _add_orientation_sections(document)
    _add_risk_sections(document, rows)
    _add_annexes(document)

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9)


def _add_cover(document: Document, context: ContractContext) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MAPA DE GERENCIAMENTO DE RISCOS")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(context.objeto or "<Nome da contratação>").bold = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{context.area_demandante} | {date.today():%d/%m/%Y}")

    document.add_paragraph()
    table = document.add_table(rows=7, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_cell_text(table.cell(0, 0), "Participantes", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    table.cell(0, 0).merge(table.cell(0, 1))
    _set_cell_text(table.cell(1, 0), "Elaboração", bold=True)
    _set_cell_text(table.cell(1, 1), "Assinatura / Data", bold=True)
    for row in table.rows:
        for cell in row.cells:
            _shade_cell(cell, "D9EAF7" if row is table.rows[0] else None)

    document.add_paragraph()
    approval = document.add_table(rows=6, cols=2)
    approval.alignment = WD_TABLE_ALIGNMENT.CENTER
    approval.style = "Table Grid"
    _set_cell_text(approval.cell(0, 0), "Aprovação", bold=True)
    _set_cell_text(approval.cell(0, 1), "Assinatura / Data", bold=True)


def _add_history(document: Document) -> None:
    _heading(document, "HISTÓRICO DE REVISÕES DO DOCUMENTO", level=1)
    table = document.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    headers = ["DATA", "VERSÃO", "DESCRIÇÃO", "AUTOR"]
    values = [f"{date.today():%d/%m/%Y}", "1.0", "GERAÇÃO PELO PROTÓTIPO MVP", ""]
    for index, header in enumerate(headers):
        _set_cell_text(table.cell(0, index), header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(table.cell(0, index), "D9EAF7")
        _set_cell_text(table.cell(1, index), values[index])


def _add_orientation_sections(document: Document) -> None:
    _heading(document, "Orientações para acesso e uso do Mapa de Riscos", level=1)
    document.add_paragraph(
        "Este documento deverá possuir a sua Classificação da Informação de acordo com as instruções "
        "contidas nos normativos de classificação e tratamento da informação vigentes."
    )
    document.add_paragraph(
        "O Mapa de Gerenciamento de Riscos de Contratação permite a identificação, avaliação e gestão "
        "dos riscos das contratações da Dataprev."
    )

    _heading(document, "Orientações para preenchimento do Mapa de Riscos", level=1)
    for item in [
        "Probabilidade: definir valor de 1 a 5 e assinalar com X o valor adotado.",
        "Impacto: definir valor de 1 a 5 e assinalar com X o valor adotado.",
        "Nível: calcular multiplicando probabilidade pelo impacto.",
        "Estratégia: indicar Mitigar, Aceitar, Compartilhar ou Evitar.",
        "Consequências, ações preventivas, ações de contingência e observações devem ser revisadas pelo responsável.",
    ]:
        document.add_paragraph(item, style=None)


def _add_risk_sections(document: Document, rows: list[MatrixRow]) -> None:
    grouped: dict[str, list[MatrixRow]] = defaultdict(list)
    for row in rows:
        if row.selecionado:
            grouped[row.categoria or "planejamento"].append(row)

    order = ["planejamento", "selecao", "gestao", "solucao", "instalacao", "cronograma"]
    for category in order:
        category_rows = grouped.get(category, [])
        if not category_rows:
            continue
        _heading(document, CATEGORY_TITLES.get(category, category), level=1)
        for index, row in enumerate(category_rows, 1):
            _add_risk_table(document, row, f"{_category_number(category)}.{index}")


def _add_risk_table(document: Document, row: MatrixRow, display_id: str) -> None:
    table = document.add_table(rows=0, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    _add_merged_row(table, ["RISCO", display_id, row.risco], [1, 1, 4], header=True)
    _add_scale(table, "PROBABILIDADE", PROBABILITY_OPTIONS, canonical_probability(row.probabilidade))
    _add_scale(table, "IMPACTO", IMPACT_OPTIONS, canonical_impact(row.impacto))
    _add_level_scale(table, row)
    _add_strategy_scale(table, row.estrategia)
    _add_list_block(table, "CONSEQUÊNCIAS", _split_items(row.consequencia), include_status=False)
    _add_list_block(
        table,
        "AÇÕES PREVENTIVAS",
        _split_items(row.acao_preventiva),
        include_status=True,
        situation="Não iniciado",
        owner=row.responsavel,
    )
    _add_list_block(
        table,
        "AÇÕES DE CONTINGÊNCIA",
        _split_items(row.acao_contingencia),
        include_status=True,
        situation="Não iniciado",
        owner=row.responsavel,
    )
    _add_merged_row(table, ["OBSERVAÇÕES"], [6], header=True)
    observation = row.justificativa or (
        f"Probabilidade {score_value(row.probabilidade)} x Impacto {score_value(row.impacto)} = "
        f"{risk_score(row.probabilidade, row.impacto)} ({row.nivel.title()})."
    )
    _add_merged_row(table, [observation], [6])

    document.add_paragraph()


def _add_scale(table, label: str, options: list[str], selected: str) -> None:
    labels = [label, *options]
    header = table.add_row()
    for index, text in enumerate(labels):
        _set_cell_text(header.cells[index], text, bold=index == 0, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(header.cells[index], "EAF3F8" if index == 0 else "F2F2F2")
    marks = table.add_row()
    _set_cell_text(marks.cells[0], "")
    for index, option in enumerate(options, start=1):
        _set_cell_text(marks.cells[index], "X" if option == selected else "", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_level_scale(table, row: MatrixRow) -> None:
    score = risk_score(row.probabilidade, row.impacto)
    selected = _level_for_score(score)
    header = table.add_row()
    _set_cell_text(header.cells[0], "NÍVEL DE RISCO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(header.cells[0], "EAF3F8")
    for index, option in enumerate(LEVEL_OPTIONS, start=1):
        _set_cell_text(header.cells[index], option, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(header.cells[index], "F2F2F2")
    _set_cell_text(header.cells[5], "")
    marks = table.add_row()
    _set_cell_text(marks.cells[0], "")
    for index, option in enumerate(LEVEL_OPTIONS, start=1):
        _set_cell_text(marks.cells[index], "X" if option == selected else "", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_strategy_scale(table, selected: str) -> None:
    normalized = (selected or "Mitigar").strip().lower()
    header = table.add_row()
    _set_cell_text(header.cells[0], "ESTRATÉGIA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(header.cells[0], "EAF3F8")
    for index, option in enumerate(STRATEGY_OPTIONS, start=1):
        _set_cell_text(header.cells[index], option, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(header.cells[index], "F2F2F2")
    _set_cell_text(header.cells[5], "")
    marks = table.add_row()
    _set_cell_text(marks.cells[0], "")
    for index, option in enumerate(STRATEGY_OPTIONS, start=1):
        _set_cell_text(marks.cells[index], "X" if option.lower() == normalized else "", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)


def _add_list_block(
    table,
    title: str,
    items: list[str],
    include_status: bool,
    situation: str = "",
    owner: str = "",
) -> None:
    if include_status:
        header = table.add_row()
        _set_cell_text(header.cells[0], "ID", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        merged = header.cells[1].merge(header.cells[3])
        _set_cell_text(merged, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(header.cells[4], "SITUAÇÃO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(header.cells[5], "RESPONSÁVEL", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for cell in header.cells:
            _shade_cell(cell, "D9EAF7")
        for index, item in enumerate(items[:3] or [""], 1):
            data = table.add_row()
            _set_cell_text(data.cells[0], str(index), align=WD_ALIGN_PARAGRAPH.CENTER)
            merged_item = data.cells[1].merge(data.cells[3])
            _set_cell_text(merged_item, item)
            _set_cell_text(data.cells[4], situation)
            _set_cell_text(data.cells[5], owner)
    else:
        header = table.add_row()
        _set_cell_text(header.cells[0], "ID", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        merged = header.cells[1].merge(header.cells[5])
        _set_cell_text(merged, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        for cell in header.cells:
            _shade_cell(cell, "D9EAF7")
        for index, item in enumerate(items[:4] or [""], 1):
            data = table.add_row()
            _set_cell_text(data.cells[0], str(index), align=WD_ALIGN_PARAGRAPH.CENTER)
            merged_item = data.cells[1].merge(data.cells[5])
            _set_cell_text(merged_item, item)


def _add_annexes(document: Document) -> None:
    _heading(document, "ANEXO 1", level=1)
    _heading(document, "Escala de Probabilidade", level=2)
    table = document.add_table(rows=3, cols=6)
    table.style = "Table Grid"
    data = [
        [
            "Aspectos Avaliativos",
            "Evento pode ocorrer apenas em circunstâncias excepcionais",
            "Evento pode ocorrer em algum momento",
            "Evento deve ocorrer em algum momento",
            "Evento provavelmente ocorre na maioria das circunstâncias",
            "Evento esperado na maioria das circunstâncias",
        ],
        ["Frequência observada/esperada", "Muito baixa (<10%)", "Baixa (>=10% <=30%)", "Média (>30% <=50%)", "Alta (>50% <=90%)", "Muito alta (>90%)"],
        ["Peso", "1", "2", "3", "4", "5"],
    ]
    _fill_table(table, data)

    _heading(document, "ANEXO 2", level=1)
    _heading(document, "Escala de Impacto", level=2)
    table = document.add_table(rows=7, cols=6)
    table.style = "Table Grid"
    data = [
        ["Impacto", "Imagem 25%", "Operacionais 25%", "Regulatórios 25%", "Financeiros 25%", "Peso"],
        ["Orientações para atribuição de pesos", "", "", "", "", ""],
        ["", "Exposição intensa", "Afeta mais de 70% a sustentação/entrega", "Determina interrupção das atividades", "Afeta mais de 25% da receita líquida anual", "5-Muito Alto"],
        ["", "Exposição significativa", "Afeta entre 50% e 70% a sustentação/entrega", "Determina ações de caráter pecuniário", "Afeta de 15% a 25% da receita líquida anual", "4-Alto"],
        ["", "Exposição temporária", "Afeta entre 30% e 50% a sustentação/entrega", "Determina ações de caráter corretivo", "Afeta de 5% a 15% da receita líquida anual", "3-Médio"],
        ["", "Exposição limitada", "Afeta entre 10% e 30% a sustentação/entrega", "Determina ações de caráter orientativo", "Afeta até 5% da receita líquida anual", "2-Baixo"],
        ["", "Sem exposição significativa", "Afeta menos de 10% a sustentação/entrega", "Pouco ou nenhum impacto", "Sem influência significativa", "1-Muito baixo"],
    ]
    _fill_table(table, data)


def _fill_table(table, data: list[list[str]]) -> None:
    for row_index, values in enumerate(data):
        for col_index, value in enumerate(values):
            _set_cell_text(table.cell(row_index, col_index), value, bold=row_index == 0, align=WD_ALIGN_PARAGRAPH.CENTER)
            if row_index == 0:
                _shade_cell(table.cell(row_index, col_index), "D9EAF7")


def _heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 10)


def _add_merged_row(table, values: list[str], spans: list[int], header: bool = False) -> None:
    row = table.add_row()
    cursor = 0
    for value, span in zip(values, spans):
        cell = row.cells[cursor]
        if span > 1:
            cell = cell.merge(row.cells[cursor + span - 1])
        _set_cell_text(cell, value, bold=header, align=WD_ALIGN_PARAGRAPH.CENTER if header else WD_ALIGN_PARAGRAPH.LEFT)
        if header:
            _shade_cell(cell, "D9EAF7")
        cursor += span


def _set_cell_text(cell, text: str, bold: bool = False, align=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text or ""))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _shade_cell(cell, fill: str | None) -> None:
    if not fill:
        return
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _split_items(text: str) -> list[str]:
    parts = [part.strip(" .") for part in (text or "").replace("\n", ";").split(";") if part.strip()]
    return parts or [text]


def _level_for_score(score: int) -> str:
    if score >= 15:
        return "15 a 25 - Crítico"
    if score >= 8:
        return "8 a 12 - Alto"
    if score >= 4:
        return "4 a 6 - Moderado"
    return "1 a 3 - Pequeno"


def _category_number(category: str) -> int:
    return {
        "planejamento": 1,
        "selecao": 2,
        "gestao": 3,
        "solucao": 4,
        "instalacao": 4,
        "cronograma": 5,
    }.get(category, 1)
